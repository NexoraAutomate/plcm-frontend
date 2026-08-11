"""
Convert Markdown workflow specs to professionally formatted Word (.docx) files.
Font: Calibri throughout. Includes headings, lists, tables, header/footer, title.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, Cm, RGBColor, Twips


# Brand / document colours
PRIMARY = RGBColor(0x1F, 0x4E, 0x79)      # dark blue headings
ACCENT = RGBColor(0x2E, 0x75, 0xB6)       # medium blue
TABLE_HEADER_BG = "1F4E79"
TABLE_ALT_ROW = "D6E3F0"
CODE_BG = "F2F2F2"
MUTED = RGBColor(0x59, 0x59, 0x59)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_run_font(run, name="Calibri", size=None, bold=None, italic=None, color=None, mono=False):
    run.font.name = "Consolas" if mono else name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    if mono:
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        rFonts.set(qn("w:cs"), "Consolas")
    else:
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)
        rFonts.set(qn("w:eastAsia"), name)
        rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=6, after=6, line=1.15, space_after_pt=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after if space_after_pt is None else space_after_pt)
    pf.line_spacing = line


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="8FAADC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def add_page_number(paragraph):
    """Add PAGE field to paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, size=9, color=MUTED)


def add_num_pages(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " NUMPAGES "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document):
    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15

    heading_sizes = {1: 22, 2: 16, 3: 13, 4: 12}
    for level, size in heading_sizes.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = PRIMARY if level <= 2 else ACCENT
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        hpf = style.paragraph_format
        hpf.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
        hpf.space_after = Pt(8 if level <= 2 else 6)
        hpf.keep_with_next = True

    # List styles
    for name in ("List Bullet", "List Number"):
        if name in styles:
            st = styles[name]
            st.font.name = "Calibri"
            st.font.size = Pt(11)
            st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def setup_header_footer(doc: Document, title: str, subtitle: str = "PLCM Workflow Specification"):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(subtitle)
    set_run_font(run, size=9, bold=True, color=PRIMARY)
    run2 = hp.add_run("  |  ")
    set_run_font(run2, size=9, color=MUTED)
    # Truncate long titles for header
    short = title if len(title) <= 70 else title[:67] + "…"
    run3 = hp.add_run(short)
    set_run_font(run3, size=9, color=MUTED)

    # Header bottom border
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1F4E79")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Top border on footer
    pPr = fp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), "8FAADC")
    pBdr.append(top)
    pPr.append(pBdr)

    r1 = fp.add_run("Confidential  ·  Page ")
    set_run_font(r1, size=9, color=MUTED)
    add_page_number(fp)
    r2 = fp.add_run(" of ")
    set_run_font(r2, size=9, color=MUTED)
    add_num_pages(fp)


INLINE_RE = re.compile(
    r"(\*\*\*[^*]+\*\*\*"          # bold+italic
    r"|\*\*[^*]+\*\*"              # bold
    r"|\*[^*]+\*"                  # italic
    r"|`[^`]+`"                    # inline code
    r"|\[[^\]]+\]\([^)]+\)"        # links
    r")"
)


def add_inline_runs(paragraph, text: str, base_size=11, base_color=None):
    """Parse inline markdown and add runs."""
    if not text:
        return
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***") and len(part) > 6:
            run = paragraph.add_run(part[3:-3])
            set_run_font(run, size=base_size, bold=True, italic=True, color=base_color or BLACK)
        elif part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True, color=base_color or BLACK)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size, italic=True, color=base_color or BLACK)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size - 0.5, mono=True, color=RGBColor(0xC0, 0x39, 0x2B))
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                label, url = m.group(1), m.group(2)
                run = paragraph.add_run(label)
                set_run_font(run, size=base_size, color=ACCENT)
                run.underline = True
                # Append URL in muted text if it's a relative path / useful
                if url.startswith("./") or url.startswith("http"):
                    run2 = paragraph.add_run(f" ({url})")
                    set_run_font(run2, size=9, color=MUTED, italic=True)
            else:
                run = paragraph.add_run(part)
                set_run_font(run, size=base_size, color=base_color or BLACK)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size, color=base_color or BLACK)


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    # |---|---| or | --- | :---: |
    cells = [c.strip() for c in s.strip("|").split("|")]
    if not cells:
        return False
    return all(re.match(r"^:?-+:?$", c) for c in cells if c != "")


def parse_table_row(line: str) -> list[str]:
    s = line.strip().rstrip("|")
    if s.startswith("|"):
        s = s[1:]
    return [c.strip() for c in s.split("|")]


def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline_runs(p, text, base_size=10, base_color=RGBColor(0xFF, 0xFF, 0xFF))
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, TABLE_HEADER_BG)
        set_cell_borders(cell, color="1F4E79", sz="4")
        set_cell_margins(cell)

    # Body rows
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, text, base_size=10)
            if r_idx % 2 == 1:
                shade_cell(cell, TABLE_ALT_ROW)
            else:
                shade_cell(cell, "FFFFFF")
            set_cell_borders(cell, color="8FAADC", sz="4")
            set_cell_margins(cell)

    # Spacer after table
    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before=2, after=8)
    return table


def add_code_block(doc: Document, lines: list[str]):
    """Render fenced code as indented monospace paragraphs with light background feel."""
    text = "\n".join(lines)
    # Use a single-cell table for visual box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    shade_cell(cell, CODE_BG)
    set_cell_borders(cell, color="BFBFBF", sz="4")
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    # Clear default para and add lines
    p0 = cell.paragraphs[0]
    p0.clear()
    first = True
    for line in lines if lines else [""]:
        if first:
            p = p0
            first = False
        else:
            p = cell.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line=1.1)
        run = p.add_run(line if line else " ")
        set_run_font(run, size=9, mono=True, color=RGBColor(0x33, 0x33, 0x33))

    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before=2, after=6)


def add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "8FAADC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_title_block(doc: Document, title: str, meta_lines: list[str]):
    """Document title + metadata block."""
    # Document type label
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(label, before=0, after=4)
    r = label.add_run("WORKFLOW SPECIFICATION")
    set_run_font(r, size=10, bold=True, color=ACCENT)

    # Main title
    tp = doc.add_paragraph()
    set_paragraph_spacing(tp, before=0, after=10)
    tr = tp.add_run(title)
    set_run_font(tr, size=24, bold=True, color=PRIMARY)

    # Thin accent line
    add_horizontal_rule(doc)

    # Meta info as a compact definition-style block
    if meta_lines:
        for ml in meta_lines:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=1, after=1)
            add_inline_runs(p, ml, base_size=10, base_color=MUTED)
        add_horizontal_rule(doc)


def convert_md_to_docx(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_styles(doc)

    # Extract H1 title and leading metadata (**Key:** value)
    title = md_path.stem.replace("-", " ").title()
    meta: list[str] = []
    body_start = 0

    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        body_start = 1
        # Skip blank lines then collect **meta** lines until --- or ##
        i = body_start
        while i < len(lines) and lines[i].strip() == "":
            i += 1
        while i < len(lines):
            s = lines[i].strip()
            if s == "---" or s.startswith("##"):
                break
            if s.startswith("**") and "**" in s[2:]:
                meta.append(s)
                i += 1
            elif s == "":
                i += 1
                # allow one blank between meta
                continue
            else:
                break
        body_start = i

    setup_header_footer(doc, title)
    add_title_block(doc, title, meta)

    i = body_start
    n = len(lines)

    # Track list numbering for nested content (Word list styles handle visual bullets)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Skip leading --- after meta
        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < n:
                i += 1  # closing fence
            add_code_block(doc, code_lines)
            continue

        # Table
        if stripped.startswith("|") and i + 1 < n and is_table_separator(lines[i + 1]):
            header = parse_table_row(stripped)
            i += 2  # skip header + separator
            rows: list[list[str]] = []
            while i < n and lines[i].strip().startswith("|") and not is_table_separator(lines[i]):
                rows.append(parse_table_row(lines[i].strip()))
                i += 1
            # Normalize column counts
            cols = len(header)
            norm_rows = []
            for row in rows:
                if len(row) < cols:
                    row = row + [""] * (cols - len(row))
                elif len(row) > cols:
                    row = row[:cols]
                norm_rows.append(row)
            add_table(doc, header, norm_rows)
            continue

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Headings
        hm = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if hm:
            level = len(hm.group(1))
            heading_text = hm.group(2).strip()
            # Skip H1 if we already used it as title
            if level == 1 and heading_text == title:
                i += 1
                continue
            p = doc.add_heading(level=level)
            # Clear default run and add with our formatting
            p.clear()
            add_inline_runs(p, heading_text, base_size={1: 22, 2: 16, 3: 13, 4: 12}[level], base_color=PRIMARY if level <= 2 else ACCENT)
            for run in p.runs:
                run.bold = True
            i += 1
            continue

        # Checkbox list
        cm = re.match(r"^[-*]\s+\[([ xX])\]\s+(.*)$", stripped)
        if cm:
            checked = cm.group(1).lower() == "x"
            content = cm.group(2)
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, before=2, after=2)
            box = "☑ " if checked else "☐ "
            run = p.add_run(box)
            set_run_font(run, size=11, color=PRIMARY if checked else MUTED)
            add_inline_runs(p, content, base_size=11)
            i += 1
            continue

        # Bullet list (possibly nested with indent)
        bm = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if bm and not stripped.startswith("```"):
            indent_spaces = len(bm.group(1).replace("\t", "    "))
            content = bm.group(2)
            # Nested bullets: use List Bullet and indent
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, before=2, after=2)
            if indent_spaces >= 2:
                p.paragraph_format.left_indent = Inches(0.25 * (indent_spaces // 2))
            add_inline_runs(p, content, base_size=11)
            i += 1
            # Consume continuation / nested under this? handled by loop
            continue

        # Numbered list
        nm = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if nm:
            indent_spaces = len(nm.group(1).replace("\t", "    "))
            content = nm.group(3)
            p = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(p, before=2, after=2)
            if indent_spaces >= 2:
                p.paragraph_format.left_indent = Inches(0.25 * (indent_spaces // 2))
            add_inline_runs(p, content, base_size=11)
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            content = re.sub(r"^>\s?", "", stripped)
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=4, after=4)
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline_runs(p, content, base_size=11, base_color=MUTED)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # Normal paragraph — may continue until blank / special
        para_parts = [stripped]
        i += 1
        while i < n:
            nxt = lines[i]
            ns = nxt.strip()
            if (
                ns == ""
                or ns.startswith("#")
                or ns.startswith("|")
                or ns.startswith("```")
                or ns == "---"
                or re.match(r"^[-*]\s+", ns)
                or re.match(r"^\d+\.\s+", ns)
                or re.match(r"^>\s?", ns)
            ):
                break
            # soft-wrapped continuation
            para_parts.append(ns)
            i += 1

        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=2, after=6)
        add_inline_runs(p, " ".join(para_parts), base_size=11)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def main():
    base = Path(__file__).resolve().parent
    out_dir = base / "docx"
    out_dir.mkdir(exist_ok=True)

    md_files = sorted(base.glob("*.md"))
    if not md_files:
        print("No .md files found.")
        sys.exit(1)

    print(f"Converting {len(md_files)} file(s) -> {out_dir}")
    for md in md_files:
        out = out_dir / (md.stem + ".docx")
        convert_md_to_docx(md, out)
        print(f"  OK  {md.name} -> docx/{out.name}")

    print("Done.")


if __name__ == "__main__":
    main()
